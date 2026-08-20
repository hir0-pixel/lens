from docx import Document
import os


def create_sample_documents():
    os.makedirs("documents", exist_ok=True)

    # Create Remote Work Policy document
    doc1 = Document()
    doc1.add_heading("Company Remote Work Policy", level=1)
    doc1.add_paragraph("Policy Number: POL-2024-0847")
    doc1.add_paragraph("Effective Date: January 15, 2024")
    doc1.add_paragraph("Department: Human Resources")
    doc1.add_heading("Section 1: Eligibility", level=2)
    doc1.add_paragraph(
        "All full-time employees who have completed their probationary period "
        "of 90 days are eligible to apply for remote work arrangements. "
        "Contractors and part-time employees must obtain written approval from "
        "their department head and the VP of Operations before requesting "
        "remote work status."
    )
    doc1.add_heading("Section 2: Equipment and Expenses", level=2)
    doc1.add_paragraph(
        "The company will provide a one-time stipend of $1,500 for home office "
        "setup. This covers ergonomic furniture, monitors, and peripherals. "
        "Employees must submit receipts within 60 days of purchase. Internet "
        "reimbursement of up to $75 per month is available upon submission of "
        "monthly invoices to the Finance department."
    )
    doc1.add_heading("Section 3: Performance Monitoring", level=2)
    doc1.add_paragraph(
        "Remote employees are required to maintain a minimum productivity score "
        "of 85% as measured by the quarterly OKR review process. Managers must "
        "conduct weekly one-on-one check-ins via video call. Failure to meet "
        "performance targets for two consecutive quarters may result in "
        "revocation of remote work privileges."
    )
    doc1.save("documents/remote_work_policy.docx")

    # Create Q3 Financial Report
    doc2 = Document()
    doc2.add_heading("Q3 2024 Financial Report", level=1)
    doc2.add_paragraph("Report ID: FIN-Q3-2024")
    doc2.add_paragraph("Prepared by: Sarah Chen, CFO")
    doc2.add_paragraph("Date: October 15, 2024")
    doc2.add_heading("Revenue Summary", level=2)
    doc2.add_paragraph(
        "Total revenue for Q3 2024 reached $47.3 million, representing a 12% "
        "increase over Q3 2023. The Enterprise Solutions division contributed "
        "$28.1 million, while the SMB segment generated $19.2 million. The APAC "
        "region showed the strongest growth at 23%, driven by new partnerships "
        "in Japan and South Korea."
    )
    doc2.add_heading("Operating Expenses", level=2)
    doc2.add_paragraph(
        "Total operating expenses were $38.7 million, with R&D spending at "
        "$15.2 million (32% of revenue). Sales and marketing expenses decreased "
        "to $12.8 million from $14.1 million in Q2, reflecting improved "
        "efficiency in the digital acquisition channels. General and "
        "administrative costs held steady at $10.7 million."
    )
    doc2.add_heading("Key Metrics", level=2)
    doc2.add_paragraph(
        "Customer acquisition cost (CAC) decreased from $847 to $723, a 14.6% "
        "improvement. Net revenue retention rate improved to 118% from 112% in "
        "the prior quarter. The Falcon Project, our next-generation analytics "
        "platform, is on track for beta release in Q1 2025 with an allocated "
        "budget of $4.2 million."
    )
    doc2.save("documents/q3_financial_report.docx")
    # Create Vendor Agreement
    doc3 = Document()
    doc3.add_heading("Master Service Agreement", level=1)
    doc3.add_paragraph(
        "Contract Number: MSA-2024-1192"
    )
    doc3.add_paragraph(
        "Between: Acme Corporation ('Client') and CloudScale Inc. ('Provider')"
    )
    doc3.add_paragraph("Effective Date: March 1, 2024")
    doc3.add_heading("Section 4.2.1: Data Processing", level=2)
    doc3.add_paragraph(
        "The Provider shall process all Client data exclusively within data "
        "centers located in the United States and European Union. Any transfer "
        "of data to facilities outside these regions requires prior written "
        "consent from the Client's Data Protection Officer. The Provider must "
        "maintain SOC 2 Type II certification and undergo annual third-party "
        "security audits."
    )
    doc3.add_heading("Section 5.1: Service Level Agreement", level=2)
    doc3.add_paragraph(
        "The Provider guarantees 99.95% uptime for all production services, "
        "measured monthly. Scheduled maintenance windows are limited to Sundays "
        "between 02:00-06:00 UTC. For each 0.1% below the guaranteed uptime, "
        "the Client receives a 5% credit on the monthly service fee, up to a "
        "maximum of 30% of the monthly fee."
    )
    doc3.add_heading("Section 7.3: Termination", level=2)
    doc3.add_paragraph(
        "Either party may terminate this agreement with 90 days written notice. "
        "Upon termination, the Provider must return or destroy all Client data "
        "within 30 days and provide written certification of data destruction. "
        "The Provider will assist with data migration to a successor provider "
        "for up to 60 days at standard rates."
    )
    doc3.save("documents/vendor_agreement.docx")

    print("Created 3 sample documents in ./documents/")


if __name__ == "__main__":
    create_sample_documents()